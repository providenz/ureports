import os
import pandas as pd
import numpy as np
from docx import Document
from django.conf import settings
from io import BytesIO
from django.core.files.base import ContentFile
import re

from search_distr.models import Month, Settlement, Distribution, Region, Person
from utils.custom_django_functions import get_or_create_object

def proccess_bool(val):
    return "+" if val else "-"




class DataLoader:
    def __init__(self, settlement, month, distributions):
        self.settlement = settlement
        self.month = month
        self.distributions = distributions
        self.distributions = self.distributions.order_by("person__name")
        template_path = os.path.join(
            settings.BASE_DIR, "search_distr", "data_loader", "template.docx"
        )
        self.template = Document(template_path)
        self.header = self.template.tables[0]
        self.table = self.template.tables[1]

    def set_header(self):
        self.header.rows[3].cells[1].text = self.settlement.region.name
        self.header.rows[4].cells[1].text = self.settlement.district
        self.header.rows[5].cells[1].text = self.settlement.community
        self.header.rows[6].cells[1].text = self.settlement.name
        self.header.rows[7].cells[1].text = f"{self.month.month_name} {self.month.year}"

    def fill_doc(self):
        if len(self.distributions) < 1:
            return None
        counter = 1
        self.table.rows[3].cells[0].text = str(counter)
        self.table.rows[3].cells[1].text = str(self.distributions[0].person.name)
        self.table.rows[3].cells[2].text = str(self.distributions[0].person.address)
        self.table.rows[3].cells[3].text = str(self.distributions[0].person.phone)
        self.table.rows[3].cells[4].text = proccess_bool(
            self.distributions[0].person.is_idp
        )
        self.table.rows[3].cells[5].text = proccess_bool(
            self.distributions[0].person.is_pwd
        )
        self.table.rows[3].cells[6].text = proccess_bool(
            self.distributions[0].person.is_returnees
        )
        self.table.rows[3].cells[7].text = str(self.distributions[0].person.age)
        self.table.rows[3].cells[8].text = str(self.distributions[0].person.gender)
        counter += 1

        for distr in self.distributions[1:]:
            row = self.table.add_row()
            row.cells[0].text = str(counter)
            row.cells[1].text = str(distr.person.name)
            row.cells[2].text = str(distr.person.address)
            row.cells[3].text = str(distr.person.phone)
            row.cells[4].text = proccess_bool(distr.person.is_idp)
            row.cells[5].text = proccess_bool(distr.person.is_pwd)
            row.cells[6].text = proccess_bool(distr.person.is_returnees)
            row.cells[7].text = str(distr.person.age)
            row.cells[8].text = str(distr.person.gender)
            counter += 1

    def get_file(self, name):
        output = BytesIO()
        self.template.save(output)
        file = ContentFile(output.getvalue(), name + ".docx")
        return file


class DataUploader:
    def __init__(self, file):
        self.df = pd.read_excel(file)
        self.counter = 0

    def fill_data(self):
        for index, row in self.df.iterrows():
            region = get_or_create_object(Region, name=row["region"])
            settlement = get_or_create_object(
                Settlement,
                name=row["settlement"],
                community=row["community"],
                district=row["district"],
                region=region,
            )
            date = pd.to_datetime(row["date"], errors="coerce")
            month = get_or_create_object(Month, month=date.month, year=date.year)
            try:
                age = int(row["age"]) if not pd.isna(row["age"]) else 0
            except:
                age = re.search(r'\d+', row["age"]).group()
            phone = (
                "+380" + str(int(row["phone"]))[-9:]
                if not pd.isna(row["phone"])
                else "-"
            )
            name = row["full_name"] if not pd.isna(row["full_name"]) else "-"
            gender = row["gender"] if not pd.isna(row["gender"]) else "-"
            address = row["address"] if not pd.isna(row["address"]) else "-"
            person = get_or_create_object(
                Person,
                name=name,
                address=address,
                age=age,
                phone=phone,
                gender=gender,
                settlement=settlement,
            )
            distribution = get_or_create_object(
                Distribution, person=person, month=month
            )
            print(f"Created distribution {distribution}")
            self.counter += 1
        return self.counter


if __name__ == "__main__":
    counter = 1
    for setl in Settlement.objects.all():
        distributions_all = Distribution.objects.filter(person__settlement=setl)
        months = Month.objects.filter(distribution__in=distributions_all)
        for month in months:
            loader = DataLoader(setl, month)
            loader.set_header()
            loader.fill_doc()
            loader.template.save(f"{counter}.docx")
            counter += 1
